#!/usr/bin/env python3
# ════════════════════════════════════════════════════════════════
# NORZEC — Criar Boletim em DOCX (Word)
# ════════════════════════════════════════════════════════════════

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# ════════════════════════════════════════════════════════════════
# CONFIGURAÇÕES
# ════════════════════════════════════════════════════════════════

FIGURES_DIR = "figures"
OUTPUT_FILE = "NORZEC_Case_Study_001_Boletim.docx"

# Cores NORZEC
MIDNIGHT_BLUE = RGBColor(6, 47, 79)      # #062F4F
TEAL = RGBColor(20, 184, 166)            # #14B8A6
CINZA = RGBColor(74, 85, 104)            # #4A5568

# ════════════════════════════════════════════════════════════════
# VERIFICAR IMAGENS
# ════════════════════════════════════════════════════════════════

print("🔍 Verificando imagens...\n")

figuras_necessarias = [
    ("01_correlation_matrix.png", "Figura 2"),
    ("02_reconstruction_main.png", "Figura 3"),
    ("03_before_after_carousel.png", "Figura 1"),
    ("04_feature_importance_carousel.png", "Figura 4")
]

figuras_existentes = {}
for filename, label in figuras_necessarias:
    filepath = os.path.join(FIGURES_DIR, filename)
    if os.path.exists(filepath):
        figuras_existentes[filename] = filepath
        print(f"✅ {filename}")
    else:
        print(f"⚠️  {filename} NÃO ENCONTRADO")

print(f"\n{len(figuras_existentes)}/4 imagens encontradas\n")

# ════════════════════════════════════════════════════════════════
# CRIAR DOCUMENTO
# ════════════════════════════════════════════════════════════════

print("📝 Criando documento DOCX...\n")

doc = Document()

# Configurar margens
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ════════════════════════════════════════════════════════════════
# CAPA
# ════════════════════════════════════════════════════════════════

# Logo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("NORZEC")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = MIDNIGHT_BLUE

# Espaço
doc.add_paragraph()

# Título
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CASE STUDY #001")
run.font.size = Pt(48)
run.font.bold = True
run.font.color.rgb = MIDNIGHT_BLUE

# Subtítulo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Reconstrução Inteligente de Dados Meteorológicos com Machine Learning")
run.font.size = Pt(20)
run.font.color.rgb = CINZA

# Espaço
doc.add_paragraph()
doc.add_paragraph()

# Metadados
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Local: Chimoio, Moçambique  |  Período: 2015–2025  |  Registro: 4.018 Dias")
run.font.size = Pt(11)
run.font.color.rgb = CINZA

# Espaço
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Mensagem principal
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Dados imperfeitos não significam decisões imperfeitas.")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = TEAL

# Quebra de página
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# SEÇÃO 1: O DESAFIO
# ════════════════════════════════════════════════════════════════

p = doc.add_heading("O Desafio", level=1)
p.style.font.color.rgb = MIDNIGHT_BLUE

p = doc.add_paragraph(
    "Séries temporais ambientais frequentemente apresentam lacunas devido a falhas de sensores, "
    "interrupções operacionais ou limitações de monitoramento. Essas falhas reduzem a qualidade "
    "de análises climáticas, estudos de risco e processos de tomada de decisão."
)

# Box Definição
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run("Definição do Estudo")
run.font.bold = True
run.font.color.rgb = MIDNIGHT_BLUE
run.font.size = Pt(12)

items = [
    "Série meteorológica diária",
    "11 anos de observações contínuas",
    "25% dos registros removidos artificialmente para simular falhas reais",
    "Objetivo: reconstruir com fidelidade suficiente para decisão"
]

for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)

# Imagem Figura 1
if "03_before_after_carousel.png" in figuras_existentes:
    doc.add_paragraph()
    try:
        doc.add_picture(figuras_existentes["03_before_after_carousel.png"], width=Inches(5.5))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "Figura 1: Série temporal de temperatura com 25% de falhas introduzidas artificialmente "
            "(ANTES, esquerda) e série reconstruída via Random Forest (DEPOIS, direita)."
        )
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = CINZA
    except Exception as e:
        print(f"⚠️  Erro ao inserir Figura 1: {e}")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# SEÇÃO 2: DADOS UTILIZADOS
# ════════════════════════════════════════════════════════════════

p = doc.add_heading("Dados Utilizados", level=1)
p.style.font.color.rgb = MIDNIGHT_BLUE

p = doc.add_paragraph(
    "O estudo utiliza nove variáveis meteorológicas obtidas da plataforma NASA POWER LARC, "
    "um repositório público de dados climáticos de alta resolução espacial."
)

# Tabela
table = doc.add_table(rows=10, cols=4)
table.style = 'Light Grid Accent 1'

# Header
header_cells = table.rows[0].cells
header_cells[0].text = "Variável"
header_cells[1].text = "Unidade"
header_cells[2].text = "Descrição"
header_cells[3].text = "Tipo"

# Colorir header
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    # Fundo
    from docx.oxml import parse_xml
    shading_elm = parse_xml(r'<w:shd {} w:fill="062F4F"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
    cell._element.get_or_add_tcPr().append(shading_elm)

# Dados
dados = [
    ("T2M", "°C", "Temperatura Média Diária", "Variável-alvo"),
    ("T2M_MAX", "°C", "Temperatura Máxima Diária", "Preditor"),
    ("T2M_MIN", "°C", "Temperatura Mínima Diária", "Preditor"),
    ("RH2M", "%", "Umidade Relativa", "Preditor"),
    ("PRECTOTCORR", "mm", "Precipitação Total Corrigida", "Preditor"),
    ("WS2M", "m/s", "Velocidade do Vento", "Preditor"),
    ("ALLSKY_SFC_SW_DWN", "W/m²", "Radiação Solar Incidente", "Preditor"),
    ("PS", "kPa", "Pressão Atmosférica ao Nível do Solo", "Preditor"),
    ("T2MDEW", "°C", "Ponto de Orvalho", "Preditor"),
]

for i, (var, unit, desc, tipo) in enumerate(dados, start=1):
    row_cells = table.rows[i].cells
    row_cells[0].text = var
    row_cells[1].text = unit
    row_cells[2].text = desc
    row_cells[3].text = tipo

doc.add_paragraph()
p = doc.add_paragraph(
    "Fonte: NASA POWER LARC (power.larc.nasa.gov) — Coordenadas: 19.1164°S, 33.4833°E — Período: 2015–2025"
)
p.paragraph_format.left_indent = Inches(0.5)
run = p.runs[0]
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = CINZA

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# SEÇÃO 3: ABORDAGEM ANALÍTICA
# ════════════════════════════════════════════════════════════════

p = doc.add_heading("Abordagem Analítica", level=1)
p.style.font.color.rgb = MIDNIGHT_BLUE

p = doc.add_paragraph(
    "O estudo segue uma metodologia estruturada em sete etapas sequenciais, "
    "desde a exploração dos dados até a validação dos resultados."
)

# Fluxograma
steps = [
    "Dados Originais",
    "Introdução de Falhas (25%)",
    "Análise Exploratória e Correlação",
    "Matriz de Correlação Física",
    "Treinamento de Random Forest",
    "Reconstrução das Lacunas",
    "Validação Estatística"
]

for step in steps:
    p = doc.add_paragraph(step)
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.runs[0]
    run.font.bold = True
    run.font.color.rgb = MIDNIGHT_BLUE

# Imagem Figura 2
if "01_correlation_matrix.png" in figuras_existentes:
    doc.add_paragraph()
    try:
        doc.add_picture(figuras_existentes["01_correlation_matrix.png"], width=Inches(5.5))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "Figura 2: Heatmap de correlação entre variáveis meteorológicas. "
            "Tons verdes indicam correlações positivas; tons avermelhados, negativas."
        )
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = CINZA
    except Exception as e:
        print(f"⚠️  Erro ao inserir Figura 2: {e}")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# SEÇÃO 4: RESULTADOS
# ════════════════════════════════════════════════════════════════

p = doc.add_heading("Resultados", level=1)
p.style.font.color.rgb = MIDNIGHT_BLUE

p = doc.add_paragraph(
    "O modelo Random Forest alcançou precisão excepcional na reconstrução de temperatura, "
    "validado em 1.004 registros de teste correspondentes aos valores removidos artificialmente."
)

# Métricas
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MAE: 0.25°C  |  RMSE: 0.33°C  |  R²: 0.9909")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = MIDNIGHT_BLUE

doc.add_paragraph()
p = doc.add_paragraph(
    "O modelo explica 99.09% da variância em temperatura, com erro médio absoluto inferior a 0.3°C — "
    "fidelidade suficiente para análises climáticas, previsão agrícola e decisões operacionais em instituições públicas e privadas."
)

# Imagem Figura 3
if "02_reconstruction_main.png" in figuras_existentes:
    doc.add_paragraph()
    try:
        doc.add_picture(figuras_existentes["02_reconstruction_main.png"], width=Inches(5.5))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "Figura 3 (Painel 2): Série de validação — valores originais (azul) sobreposto a valores reconstruídos "
            "via Random Forest (verde). O traçado paralelo confirma alta fidelidade."
        )
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = CINZA
    except Exception as e:
        print(f"⚠️  Erro ao inserir Figura 3: {e}")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# SEÇÃO 5: O QUE O MODELO APRENDEU
# ════════════════════════════════════════════════════════════════

p = doc.add_heading("O Que o Modelo Aprendeu", level=1)
p.style.font.color.rgb = MIDNIGHT_BLUE

p = doc.add_paragraph(
    "A análise de importância das variáveis revela quais preditores foram mais relevantes "
    "para a reconstrução de temperatura, validando a coerência física do modelo."
)

# Imagem Figura 4
if "04_feature_importance_carousel.png" in figuras_existentes:
    doc.add_paragraph()
    try:
        doc.add_picture(figuras_existentes["04_feature_importance_carousel.png"], width=Inches(5.5))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "Figura 4: Importância das variáveis (Ganho de Gini). "
            "Temperatura mínima (T2M_MIN) e máxima (T2M_MAX) dominam, respondendo por ~98% da capacidade preditiva."
        )
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = CINZA
    except Exception as e:
        print(f"⚠️  Erro ao inserir Figura 4: {e}")

# Descoberta Principal
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Descoberta Principal: ")
run.font.bold = True
run.font.color.rgb = MIDNIGHT_BLUE
run = p.add_run(
    "Temperatura mínima e máxima foram responsáveis por aproximadamente 98% da capacidade preditiva do modelo. "
    "Esse resultado demonstra forte coerência física entre as variáveis: por definição termodinâmica, "
    "a temperatura média é determinada principalmente pelos extremos diários em climas tropicais."
)

# Implicação Técnica
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Implicação Técnica: ")
run.font.bold = True
run.font.color.rgb = MIDNIGHT_BLUE
run = p.add_run(
    "A concentração de importância em poucas variáveis sugere que monitores meteorológicos podem priorizar "
    "essas medições para garantir qualidade dos dados."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# SEÇÃO 6: IMPACTO OPERACIONAL
# ════════════════════════════════════════════════════════════════

p = doc.add_heading("Impacto Operacional", level=1)
p.style.font.color.rgb = MIDNIGHT_BLUE

p = doc.add_paragraph(
    "A capacidade de reconstrução de dados meteorológicos com alta fidelidade oferece "
    "aplicações práticas em múltiplos setores."
)

# Tabela de impacto
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'

impactos = [
    ("Meteorologia", "Reconstrução de séries históricas, preenchimento de gaps em observatórios"),
    ("Agricultura", "Zoneamento agroclimático, cálculo de índices de déficit hídrico"),
    ("Energia", "Completação de bases de geração solar/eólica"),
    ("Recursos Hídricos", "Séries de precipitação para modelagem hidrológica"),
    ("Saúde Pública", "Estudos ambientais e epidemiológicos (malária, dengue)"),
    ("Infraestrutura", "Monitoramento climático para planejamento urbano"),
]

header_cells = table.rows[0].cells
header_cells[0].text = "Setor"
header_cells[1].text = "Aplicação"

for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    from docx.oxml import parse_xml
    shading_elm = parse_xml(r'<w:shd {} w:fill="14B8A6"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
    cell._element.get_or_add_tcPr().append(shading_elm)

for i, (setor, app) in enumerate(impactos, start=1):
    row_cells = table.rows[i].cells
    row_cells[0].text = setor
    row_cells[1].text = app

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# SEÇÃO 7: POR QUE ISSO IMPORTA
# ════════════════════════════════════════════════════════════════

p = doc.add_heading("Por Que Isso Importa", level=1)
p.style.font.color.rgb = MIDNIGHT_BLUE

p = doc.add_paragraph(
    "Muitos projetos em África operam com bases de dados incompletas. "
    "A capacidade de recuperar informações ausentes aumenta a confiabilidade das análises, "
    "reduz custos associados à recolha de novos dados e acelera processos de tomada de decisão."
)

doc.add_paragraph()
p = doc.add_paragraph(
    "Em contextos de recursos limitados, dados meteorológicos lacunosos são uma barreira recorrente à pesquisa. "
    "Este estudo demonstra que abordagens de Machine Learning pragmáticas e validadas podem contornar essa limitação."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# SEÇÃO 8: CAPACIDADES NORZEC
# ════════════════════════════════════════════════════════════════

p = doc.add_heading("Da Ciência dos Dados à Tomada de Decisão", level=1)
p.style.font.color.rgb = MIDNIGHT_BLUE

p = doc.add_paragraph(
    "NORZEC oferece expertise integrada em ciência de dados, engenharia ambiental "
    "e inteligência climatológica para transformar dados em decisões acionáveis."
)

capacidades = [
    "Data Science",
    "Machine Learning",
    "Climate Intelligence",
    "Geospatial Analytics",
    "Environmental Monitoring",
    "Risk Intelligence",
    "Decision Support Systems",
    "Research & Development"
]

for cap in capacidades:
    p = doc.add_paragraph(cap, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# PÁGINA FINAL
# ════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Transformamos dados dispersos em inteligência acionável.")
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = MIDNIGHT_BLUE

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("NORZEC")
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = MIDNIGHT_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Engineering • Data • Innovation")
run.font.size = Pt(11)
run.font.color.rgb = TEAL

doc.add_paragraph()
doc.add_paragraph()

# Contatos
contatos = [
    ("Website", "norzec.github.io"),
    ("Email", "norzecmz@gmail.com"),
    ("WhatsApp", "+258 87 878 7963"),
    ("LinkedIn", "linkedin.com/company/norzec-lda")
]

for label, valor in contatos:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{label}: ")
    run.font.bold = True
    run.font.color.rgb = MIDNIGHT_BLUE
    run = p.add_run(valor)
    run.font.color.rgb = CINZA

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Boletim Técnico — Junho 2026")
run.font.size = Pt(9)
run.font.color.rgb = CINZA

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("© NORZEC, LDA. Todos os direitos reservados.")
run.font.size = Pt(9)
run.font.color.rgb = CINZA

# ════════════════════════════════════════════════════════════════
# SALVAR
# ════════════════════════════════════════════════════════════════

print(f"💾 Salvando documento...\n")

try:
    doc.save(OUTPUT_FILE)
    print(f"✅ Documento criado: {OUTPUT_FILE}\n")
    print("=" * 60)
    print("✨ SUCESSO!")
    print("=" * 60)
    print(f"""
Arquivo: {OUTPUT_FILE}

Como usar:
1. Abra em Microsoft Word (ou LibreOffice)
2. Edite o que precisar
3. File → Export as PDF
4. Pronto para distribuição!

Próximos passos:
[ ] Editar/revisar no Word (se necessário)
[ ] Exportar para PDF
[ ] Publicar no site/email/WhatsApp

Status: PRONTO PARA SEGUNDA-FEIRA! 🚀
""")
except Exception as e:
    print(f"❌ Erro ao salvar: {e}\n")